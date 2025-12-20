import asyncio
import sys
import time
from io import BytesIO
from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeout
import random
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.shared import Cm
import os
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from threading import Lock
from typing import Optional, Tuple, List, Union
import pyperclip

LIST_OF_USER_AGENTS = [
    'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
    'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36',
    'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
    'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
]
import sys
import asyncio

# Fix for Windows + Playwright + Streamlit compatibility
if sys.platform == 'win32':
    asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

WEBSITE_URL = "https://texttohuman.com"
# Thread-safe print lock
print_lock = Lock()

def thread_safe_print(*args, **kwargs):
    """Thread-safe print function"""
    with print_lock:
        print(*args, **kwargs)

def get_random_user_agent():
    return random.choice(LIST_OF_USER_AGENTS)

def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    """
    if hasattr(parent, 'paragraphs'):
        for paragraph in parent.paragraphs:
            yield paragraph
    if hasattr(parent, 'tables'):
        for table in parent.tables:
            yield table

def extract_text_and_runs(doc_path: str) -> Tuple[Document, List[Tuple[Union[Paragraph, _Cell], str]]]:
    """
    Extracts text from a DOCX document while preserving the original run/paragraph objects.
    
    Returns:
        Tuple[Document, List[Tuple[Union[Paragraph, _Cell], str]]]: The document object and a list of 
        (object, text) tuples for all text-containing blocks.
    """
    doc = Document(doc_path)
    text_blocks = []
    
    # Iterate through all blocks (paragraphs and tables) in the document body
    for block in iter_block_items(doc.body):
        if isinstance(block, Paragraph):
            if block.text.strip():
                text_blocks.append((block, block.text))
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    # Iterate through all blocks (paragraphs) within the cell
                    for cell_block in iter_block_items(cell):
                        if isinstance(cell_block, Paragraph):
                            if cell_block.text.strip():
                                text_blocks.append((cell_block, cell_block.text))
    
    return doc, text_blocks

def replace_text_in_paragraph(paragraph: Paragraph, new_text: str):
    """
    Replaces the text in a paragraph while preserving the formatting of the first run.
    """
    # Clear existing runs
    for run in paragraph.runs:
        run.clear()
    
    # Add the new text to the first run, or create a new run if none existed
    if paragraph.runs:
        first_run = paragraph.runs[0]
    else:
        first_run = paragraph.add_run()
        
    first_run.text = new_text

def read_docx_and_humanize(file_path: str, page, chunk_size: int = 2000) -> Optional[BytesIO]:
    """
    Reads a DOCX, humanizes the text content element by element, and returns 
    the modified DOCX as a BytesIO object.
    """
    try:
        doc, text_blocks = extract_text_and_runs(file_path)
        
        if not text_blocks:
            thread_safe_print("No text found in the document to humanize.")
            return None
            
        thread_safe_print(f"Found {len(text_blocks)} text blocks to process.")
        
        # Prepare chunks for humanization (based on text blocks)
        text_to_humanize = [text for _, text in text_blocks]
        
        # Simple chunking for the web service, keeping track of original block indices
        chunks = []
        current_chunk_text = ""
        current_chunk_indices = []
        
        for i, text in enumerate(text_to_humanize):
            # Estimate word count (simple split)
            text_word_count = len(text.split())
            current_word_count = len(current_chunk_text.split())
            
            if current_word_count + text_word_count > chunk_size and current_chunk_text:
                chunks.append({
                    'text': current_chunk_text.strip(),
                    'indices': current_chunk_indices
                })
                current_chunk_text = text + "\n\n"
                current_chunk_indices = [i]
            else:
                current_chunk_text += text + "\n\n"
                current_chunk_indices.append(i)
        
        if current_chunk_text.strip():
            chunks.append({
                'text': current_chunk_text.strip(),
                'indices': current_chunk_indices
            })
            
        thread_safe_print(f"Split into {len(chunks)} chunks for web service.")
        
        humanized_texts = {} # {original_block_index: humanized_text}
        
        for i, chunk_data in enumerate(chunks):
            thread_safe_print(f"\n{'='*70}")
            thread_safe_print(f"Processing Chunk {i+1}/{len(chunks)}: {len(chunk_data['text'].split())} words")
            thread_safe_print(f"{'='*70}")
            
            # Humanize the chunk
            humanized_chunk_text = get_texttohuman_humanizer_final(chunk_data['text'], page, save_debug=False)
            
            if humanized_chunk_text:
                # Simple split of the humanized chunk back into blocks. 
                # This is a simplification and assumes the humanizer preserves the number of paragraphs.
                # A more robust solution would require a more sophisticated text alignment algorithm.
                # For now, we rely on the fact that the humanizer output is a single block of text.
                
                # Since the humanizer returns a single block of text, we'll replace the entire chunk's content
                # with the humanized text, and then try to re-split it based on the original block count.
                
                # The safest bet is to replace the text in the first block of the chunk and clear the rest,
                # but that loses content. The best we can do with the current web service is to 
                # assume the humanized text is a single block and replace the text of the first block.
                # However, the original code used a simple \n\n split, so we'll try to mimic that.
                
                # For a single chunk, we assume the humanized text corresponds to the original blocks
                # joined by \n\n.
                
                # Split the humanized text back into blocks based on the separator used for joining.
                # This is highly fragile, but necessary given the current architecture.
                humanized_blocks = humanized_chunk_text.split('\n\n')
                
                # Pad or truncate the humanized blocks to match the original block count
                original_block_count = len(chunk_data['indices'])
                
                if len(humanized_blocks) < original_block_count:
                    # Pad with empty strings if the humanizer merged blocks
                    humanized_blocks.extend([''] * (original_block_count - len(humanized_blocks)))
                elif len(humanized_blocks) > original_block_count:
                    # Truncate or merge extra blocks if the humanizer split blocks
                    # For simplicity, we'll truncate the extra blocks
                    humanized_blocks = humanized_blocks[:original_block_count]
                
                # Map the humanized text back to the original block indices
                for j, original_index in enumerate(chunk_data['indices']):
                    humanized_texts[original_index] = humanized_blocks[j]
            else:
                thread_safe_print(f"✗ Chunk {i+1} returned no result. Skipping replacement for this chunk.")
        
        # Replace text in the original document structure
        for i, (block, original_text) in enumerate(text_blocks):
            if i in humanized_texts:
                new_text = humanized_texts[i]
                if isinstance(block, Paragraph):
                    replace_text_in_paragraph(block, new_text)
                elif isinstance(block, _Cell):
                    # For cells, we assume the text block was the first paragraph in the cell
                    if block.paragraphs and block.paragraphs[0].text == original_text:
                        replace_text_in_paragraph(block.paragraphs[0], new_text)
                    else:
                        # Fallback: clear cell and add new paragraph
                        for p in block.paragraphs:
                            p.clear()
                        block.text = new_text
        
        # Save to BytesIO buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
        
    except Exception as e:
        thread_safe_print(f"Error processing DOCX for humanization: {e}")
        return None

# The original read_docx_with_spacing is no longer needed for the main flow, 
# but we keep it for compatibility if other parts of the code still use it.
def read_docx_with_spacing(file_path):
    """
    Read a DOCX file and return text while maintaining spacing and formatting.
    
    Args:
        file_path: str - Path to the DOCX file
        
    Returns:
        str: The extracted text with preserved spacing and line breaks
    """
    try:
        # Load the document
        doc = Document(file_path)
        
        # List to store all text elements
        text_elements = []
        
        # Iterate through all paragraphs
        for paragraph in doc.paragraphs:
            # Get the paragraph text
            para_text = paragraph.text
            
            # Preserve empty lines (paragraphs with no text)
            if not para_text.strip():
                text_elements.append('')
            else:
                text_elements.append(para_text)
        
        # Join all elements with newlines to maintain structure
        full_text = '\n'.join(text_elements)
        
        return full_text
        
    except ImportError:
        print("Error: python-docx library not installed. Install it using: pip install python-docx")
        return None
    except FileNotFoundError:
        print(f"Error: File not found at path: {file_path}")
        return None
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
        return None

class PlaywrightHumanizer:
    """Context manager for Playwright browser instance"""
    
    def __init__(self, headless=True, debug=False):
        self.playwright = None
        self.browser = None
        self.context = None
        self.page = None
        self.headless = headless
        self.debug = debug
    
    def __enter__(self):
        try:
            self.playwright = sync_playwright().start()
            
            # Launch browser with options
            self.browser = self.playwright.chromium.launch(
                headless=self.headless,
                args=[
                    '--no-sandbox',
                    '--disable-blink-features=AutomationControlled',
                    '--disable-dev-shm-usage'
                ]
            )
        except Exception as e:
            if "Executable doesn't exist" in str(e):
                print("\n" + "="*70)
                print("ERROR: Playwright browsers are not installed!")
                print("="*70)
                print("\nPlease run the following command to install browsers:")
                print("\n    playwright install chromium")
                print("\nOr install all browsers with:")
                print("\n    playwright install")
                print("\n" + "="*70 + "\n")
                raise SystemExit(1)
            else:
                raise
        
        # Create context with custom user agent and permissions
        self.context = self.browser.new_context(
            user_agent=get_random_user_agent(),
            viewport={'width': 1920, 'height': 1080},
            permissions=['clipboard-read', 'clipboard-write']
        )
        
        # Enable debug mode if requested
        if self.debug:
            self.context.set_default_timeout(120000)  # 2 minutes for debug
        
        # Create page
        self.page = self.context.new_page()
        self.page.set_default_timeout(60000)  # 60 seconds
        
        # Navigate to website
        print(f"Navigating to {WEBSITE_URL}...")
        self.page.goto(WEBSITE_URL, wait_until='networkidle')
        print("Page loaded successfully!")
        
        # Take screenshot if debug mode
        if self.debug:
            self.page.screenshot(path="debug_page_loaded.png")
            print("Screenshot saved: debug_page_loaded.png")
        
        return self.page
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        if self.page:
            self.page.close()
        if self.context:
            self.context.close()
        if self.browser:
            self.browser.close()
        if self.playwright:
            self.playwright.stop()

def get_huminizer_chrome_driver():
    """
    Create and return a Playwright page instance.
    Note: Use context manager PlaywrightHumanizer instead for proper resource management.
    """
    playwright = sync_playwright().start()
    browser = playwright.chromium.launch(
        headless=True,
        args=[
            '--no-sandbox',
            '--disable-blink-features=AutomationControlled',
            '--disable-dev-shm-usage'
        ]
    )
    
    context = browser.new_context(
        user_agent=get_random_user_agent(),
        viewport={'width': 1920, 'height': 1080},
        permissions=['clipboard-read', 'clipboard-write']
    )
    
    page = context.new_page()
    page.set_default_timeout(60000)
    page.goto(WEBSITE_URL, wait_until='networkidle')
    
    # Store references for cleanup
    page._playwright = playwright
    page._browser = browser
    page._context = context
    
    return page

def split_text_preserve_paragraphs_and_newlines(text, chunk_size=2000):
    """
    Split text into chunks while preserving paragraph boundaries and all newlines.
    
    Args:
        text: str - The text to split
        chunk_size: int - Target number of words per chunk (default: 2000)
        
    Returns:
        list: List of text chunks with preserved formatting
    """
    lines = text.split('\n')
    
    chunks = []
    current_chunk = []
    current_word_count = 0
    
    for i, line in enumerate(lines):
        line_words = line.split()
        line_word_count = len(line_words)
        
        if current_word_count + line_word_count > chunk_size and current_chunk:
            chunks.append('\n'.join(current_chunk))
            current_chunk = [line]
            current_word_count = line_word_count
        else:
            current_chunk.append(line)
            current_word_count += line_word_count
    
    if current_chunk:
        chunks.append('\n'.join(current_chunk))
    
    return chunks

def get_Zero_Human_Alternative(dialog, page):
    """
    Get the alternative button with "Human" type and 0% score.
    Retries up to 6 times by clicking reload if not found.
    
    Args:
        dialog: Locator - The dialog containing alternatives
        page: Page - Playwright page instance
        
    Returns:
        str: The text of the best alternative, or None if not found
    """
    max_retries = 6
    
    for attempt in range(max_retries):
        print(f"   Attempt {attempt + 1}/{max_retries} to find 0% Human alternative...")
        
        try:
            # Get alternatives container
            alternatives_container = dialog.locator('div.space-y-2').first
            alternatives_container.wait_for(state='visible', timeout=30000)
            
            alternative_buttons = alternatives_container.locator('button').all()
            
            if not alternative_buttons:
                print(f"   ✗ No alternative buttons found on attempt {attempt + 1}")
            else:
                # Process each button to find 0% Human alternative
                for button in alternative_buttons:
                    try:
                        spans_container = button.locator('div.flex.items-center.gap-2.text-xs').first
                        spans = spans_container.locator('span').all()
                        
                        if len(spans) >= 2:
                            alternative_type = spans[0].inner_text()
                            alternative_score_text = spans[1].inner_text()
                            
                            if alternative_type == "Human":
                                try:
                                    alternative_score = float(alternative_score_text.replace('%', ''))
                                except ValueError:
                                    print(f"   ⚠ Could not parse score: {alternative_score_text}")
                                    continue
                                
                                alternative_text_elem = button.locator('p.text-sm.text-foreground.flex-1').first
                                alternative_text = alternative_text_elem.inner_text()
                                
                                print(f"   Found Human alternative: {alternative_score}% - {alternative_text[:50]}...")
                                
                                if alternative_score < 15.0:
                                    print(f"   ✓ Found 0% Human alternative!")
                                    button.click()
                                    return alternative_text
                        
                    except Exception as e:
                        print(f"   ⚠ Error processing button: {e}")
                        continue
            
            # If not found and not the last attempt, try reloading
            if attempt < max_retries - 1:
                try:
                    reload_container = dialog.locator('div.flex.justify-end').first
                    reload_button = reload_container.locator('button').first
                    
                    reload_button.click()
                    print(f"   ✓ Clicked reload button, waiting...")
                    time.sleep(2)
                    
                    # Wait for alternatives to reload
                    dialog.locator('div.space-y-2').first.wait_for(state='visible', timeout=30000)
                    
                except Exception as e:
                    print(f"   ✗ Failed to reload alternatives: {e}")
                    break
            else:
                print(f"   ✗ Max retries reached, no 0% Human alternative found")
        
        except Exception as e:
            print(f"   ✗ Error on attempt {attempt + 1}: {e}")
            if attempt < max_retries - 1:
                time.sleep(2)
            continue
    
    return None

def get_texttohuman_humanizer_final(humanize_text, page, timeout=30000, save_debug=False):
    """
    Humanize text using Playwright
    
    Args:
        humanize_text: str - Text to humanize
        page: Page - Playwright page instance
        timeout: int - Timeout in milliseconds
        save_debug: bool - Save debug screenshots on error
    """
    processing_timeout = 60
    
    try:
        print(f"Processing text with {len(humanize_text)} characters...")
        
        # Wait for page to be fully loaded
        page.wait_for_load_state('networkidle', timeout=timeout)
        time.sleep(2)
        
        # Wait for textarea and clear it
        print("Locating textarea...")
        textarea = page.locator('textarea[data-slot="textarea"]').first
        textarea.wait_for(state='visible', timeout=timeout)
        
        # Clear and focus textarea
        textarea.click()
        textarea.fill('')
        time.sleep(1)
        
        # Scroll textarea into view
        textarea.scroll_into_view_if_needed()
        
        # Try multiple methods to input text
        print("Attempting to paste text...")
        
        # Method 1: Try using clipboard paste button
        try:
            print("Trying direct input method...")
            
            # Method 2: Direct fill
            textarea.fill(humanize_text)
            time.sleep(1)
            
            # Method 3: Type with keyboard simulation (fallback)
            if not textarea.input_value():
                print("Direct fill failed, trying keyboard input...")
                textarea.click()
                page.keyboard.insert_text(humanize_text)
                time.sleep(1)
            
        except Exception as e:
            print(f"Paste button method failed: {e}")
            pyperclip.copy(humanize_text)
            paste_button = page.locator('button.bg-primary\\/10').first
            
            if paste_button.is_visible(timeout=5000):
                print("Found paste button, clicking...")
                paste_button.click()
                time.sleep(2)
            else:
                raise Exception("Paste button not visible")
        
        # Verify text was entered
        current_value = textarea.input_value()
        print(f"Textarea now has {len(current_value)} characters")
        
        if len(current_value) < 10:
            if save_debug:
                page.screenshot(path="debug_text_input_failed.png")
            raise Exception("Failed to enter text into textarea")
        
        # Wait for and click humanize button - try multiple selectors
        print("Looking for Humanize button...")
        
        humanize_button = None
        button_selectors = [
            'button[data-slot="button"]:not([disabled])',
            'button:has-text("Humanize")',
            'button:has-text("Humanize Now")',
            'button.inline-flex:not([disabled])',
        ]
        
        humanize_button = page.get_by_role("button", name="Humanize Now")
        print("Found:", humanize_button.count())
        print("Visible:", humanize_button.is_visible())
        print("Enabled:", humanize_button.is_enabled())

        
        if humanize_button is None:
            # Debug: Print all buttons on page
            print("Could not find humanize button. Available buttons:")
            all_buttons = page.locator('button').all()
            for idx, btn in enumerate(all_buttons[:10]):  # Show first 10 buttons
                try:
                    btn_text = btn.inner_text()
                    if btn_text.strip() == "Humanize Now":
                        humanize_button = btn
                    btn_disabled = btn.get_attribute('disabled')
                    print(f"  Button {idx}: '{btn_text}' (disabled={btn_disabled})")
                except:
                    pass
            
            if save_debug:
                page.screenshot(path="debug_button_not_found.png")
            
            raise Exception("Could not locate Humanize button")
        
        # Click the humanize button
        print("Clicking Humanize button...")
        humanize_button.click()
        
        # Monitor processing status
        start_time = time.time()
        max_wait_time = processing_timeout
        check_interval = 2
        last_status = ""
        
        while True:
            elapsed_time = time.time() - start_time
            
            if elapsed_time > max_wait_time:
                thread_safe_print(f"Timeout after {elapsed_time:.1f} seconds")
                break
            
            try:
                status_div = page.locator('div.flex.items-center.gap-4.text-xs.text-primary').first
                if status_div.is_visible():
                    status_text = status_div.inner_text().strip()
                    
                    if status_text and status_text != last_status:
                        thread_safe_print(f"⚡ Autopilot: {status_text} ({int(elapsed_time)}s elapsed)")
                        last_status = status_text
            except:
                pass
            
            try:
                output_element = page.locator('div.p-4.overflow-y-auto.rounded-lg.h-full.text-foreground.bg-background').first
                if output_element.is_visible() and output_element.inner_text().strip():
                    break
            except:
                pass
            
            time.sleep(check_interval)
        
        # Get output text
        output_element = page.locator('div.p-4.overflow-y-auto.rounded-lg.h-full.text-foreground.bg-background').first
        output_element.wait_for(state='visible', timeout=timeout)
        
        
        humanized_text = output_element.inner_text()
        print(humanized_text)
        humanize_text1 = humanized_text
        
        # Process marks (highlighted sections)
        marks = output_element.locator('mark').all()
        
        if marks:
            for i, mark in enumerate(marks):
                mark_class = mark.get_attribute('class') or ""
                
                if ('bg-yellow-100' in mark_class) or ('bg-yellow-900' in mark_class) or \
                   ('bg-red-100' in mark_class) or ('bg-red-900' in mark_class):
                    
                    mark_type = "yellow" if 'yellow' in mark_class else "red"
                    print(f"\n🔄 Processing {mark_type} mark {i+1}/{len(marks)}")
                    
                    mark_text = mark.inner_text()
                    print(f"   Original text: {mark_text[:80]}...")
                    
                    try:
                        mark.scroll_into_view_if_needed()
                        time.sleep(1)
                        mark.click()
                        
                        # Wait for dialog
                        dialog = page.locator('div[role="dialog"]').first
                        dialog.wait_for(state='visible', timeout=30000)
                        
                        # Wait for alternatives to load
                        dialog.locator('div.space-y-2').first.wait_for(state='visible', timeout=30000)
                        print("   ✓ Dialog loaded with alternatives")
                        
                        # If mark_text is empty, get from textarea
                        if mark_text.strip() == "":
                            try:
                                textarea_in_dialog = dialog.locator('textarea').first
                                mark_text = textarea_in_dialog.input_value()
                                print(f"   Retrieved text from textarea: {mark_text[:80]}...")
                            except Exception as e:
                                print(f"   ✗ Failed to get textarea text: {e}")
                                continue
                        
                        # Get best alternative
                        best_alternative_text = get_Zero_Human_Alternative(dialog, page)
                        
                        if best_alternative_text is not None:
                            print(f"   ✓ Best alternative text: {best_alternative_text[:80]}...")
                            humanize_text1 = humanize_text1.replace(mark_text, best_alternative_text, 1)
                            print(f"   ✓ Replaced text in humanize_text1")
                        else:
                            print("   ✗ No 0% Human alternative found after all retries")
                        
                        # Close dialog
                        # try:
                        #     if dialog.is_visible():
                        #         close_button = dialog.locator('button[data-slot="dialog-close"]').first
                        #         close_button.click()
                        #         time.sleep(1)
                        # except Exception as e:
                        #     print(f"   ⚠ Failed to close dialog: {e}")
                            
                    except Exception as e:
                        print(f"   ✗ Failed to process mark: {e}")
                        continue
        
        # Get output text
        
        
        output_element1 = page.locator('div.p-4.overflow-y-auto.rounded-lg.h-full.text-foreground.bg-background').first
        output_element1.wait_for(state='visible', timeout=timeout)
        
        humanized_text_final = output_element1.inner_text()
        thread_safe_print("Final Output element is visible, retrieving text...")
        thread_safe_print("" + humanized_text_final)
        # Save to file
        with open("humanized_text_final.txt", "w", encoding="utf-8") as f:
            f.write(humanized_text_final)
            
        with open("humanized_text.txt", "w", encoding="utf-8") as f:
            f.write(humanize_text1)
        
        return humanize_text1
    
    except Exception as e:
        print(f"Error occurred: {e}")
        return None

if __name__ == "__main__":
    docx_file = r"Manual Introduction.docx"
    
    # Enable debug mode: headless=False to see browser, debug=True for screenshots
    DEBUG_MODE = True  # Set to True to see what's happening
    
    with PlaywrightHumanizer(headless=not DEBUG_MODE, debug=DEBUG_MODE) as page:
        # Use the new function for DOCX processing
        docx_buffer = read_docx_and_humanize(docx_file, page, chunk_size=2000)
        
        if docx_buffer:
            output_path = "humanized_output.docx"
            with open(output_path, "wb") as f:
                f.write(docx_buffer.getbuffer())
            
            print("\n" + "="*70)
            print("✓ Processing Complete!")
            print(f"✓ Output saved to: {output_path}")
            print("="*70)
        else:
            print("\n" + "="*70)
            print("✗ No DOCX was generated.")
            print("="*70)
