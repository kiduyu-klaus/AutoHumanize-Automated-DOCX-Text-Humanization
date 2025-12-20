import streamlit as st
import time
import subprocess
import os
import sys
from pathlib import Path
import tempfile
from docx import Document
from docx.shared import Pt
from io import BytesIO
from datetime import datetime

# Function to check and install Playwright browsers
def ensure_playwright_installed():
    """Check if Playwright browsers are installed, install if not"""
    cache_dir = Path.home() / ".cache" / "ms-playwright"
    
    # Check if chromium is installed
    chromium_installed = False
    if cache_dir.exists():
        for item in cache_dir.iterdir():
            if item.is_dir() and "chromium" in item.name.lower():
                chromium_installed = True
                break
    
    if not chromium_installed:
        st.info("🎭 Installing Playwright browsers for the first time... This may take a few minutes.")
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        try:
            status_text.text("Installing Playwright Chromium browser...")
            progress_bar.progress(30)
            
            # Install chromium browser only
            result = subprocess.run(
                [sys.executable, "-m", "playwright", "install", "chromium"],
                capture_output=True,
                text=True,
                timeout=300  # 5 minutes timeout
            )
            
            progress_bar.progress(70)
            
            if result.returncode != 0:
                st.error(f"❌ Failed to install Playwright browsers: {result.stderr}")
                st.info("💡 You may need to manually run: `playwright install chromium`")
                return False
            
            status_text.text("Installing system dependencies...")
            progress_bar.progress(90)
            
            # Try to install dependencies (may fail without sudo, but browser will still work in most cases)
            subprocess.run(
                [sys.executable, "-m", "playwright", "install-deps", "chromium"],
                capture_output=True,
                text=True,
                timeout=120
            )
            
            progress_bar.progress(100)
            status_text.text("✅ Playwright installation complete!")
            time.sleep(2)
            progress_bar.empty()
            status_text.empty()
            
            return True
            
        except subprocess.TimeoutExpired:
            st.error("❌ Installation timed out. Please try again.")
            return False
        except Exception as e:
            st.error(f"❌ Error during installation: {str(e)}")
            st.info("💡 You may need to manually run: `playwright install chromium`")
            return False
    
    return True

# Check and install Playwright before importing the module
if 'playwright_checked' not in st.session_state:
    st.session_state.playwright_checked = False

if not st.session_state.playwright_checked:
    ensure_playwright_installed()
    st.session_state.playwright_checked = True

# Now import the texttohuman module
from texttohuman import (
    get_huminizer_chrome_driver,
    get_texttohuman_humanizer_final,
    read_docx_with_spacing, # Kept for compatibility, though not used in new DOCX flow
    split_text_preserve_paragraphs_and_newlines,
    read_docx_and_humanize # New function for DOCX processing
)

# Page configuration
st.set_page_config(
    page_title="autohumanize-app : AI Text Humanizer",
    page_icon="✨",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize session state
if 'humanized_text' not in st.session_state:
    st.session_state.humanized_text = ""
if 'processing' not in st.session_state:
    st.session_state.processing = False
if 'driver' not in st.session_state:
    st.session_state.driver = None
if 'output_filename' not in st.session_state:
    st.session_state.output_filename = ""
if 'input_filename' not in st.session_state:
    st.session_state.input_filename = ""
if 'docx_buffer' not in st.session_state:
    st.session_state.docx_buffer = None
if 'input_method' not in st.session_state:
    st.session_state.input_method = "Type/Paste Text"
if 'dark_mode' not in st.session_state:
    st.session_state.dark_mode = False # Default to light mode

# --- Custom CSS for Modern UI and Dark Mode ---
def get_custom_css(is_dark_mode):
    """Returns the custom CSS string based on the theme."""
    
    # Define colors based on theme
    if is_dark_mode:
        primary_color = "#818cf8" # Indigo 400
        secondary_color = "#10b981" # Emerald 500
        background_color = "#111827" # Gray 900
        card_bg = "#1f2937" # Gray 800
        text_color = "#f9fafb" # Gray 50
        header_gradient_start = "#a5b4fc" # Indigo 300
        header_gradient_end = "#c4b5fd" # Violet 300
        border_color = "#374151" # Gray 700
        info_bg = "#1e3a8a" # Blue 900
        info_border = "#3b82f6" # Blue 500
    else:
        primary_color = "#4f46e5" # Indigo 600
        secondary_color = "#059669" # Emerald 600
        background_color = "#f9fafb" # Gray 50
        card_bg = "white"
        text_color = "#1f2937" # Gray 800
        header_gradient_start = "#4f46e5" # Indigo 600
        header_gradient_end = "#7c3aed" # Violet 600
        border_color = "#e5e7eb" # Gray 200
        info_bg = "#eef2ff" # Indigo 50
        info_border = "#4f46e5" # Indigo 600

    css = f"""
    <style>
    /* Global Variables */
    :root {{
        --primary-color: {primary_color};
        --secondary-color: {secondary_color};
        --background-color: {background_color};
        --card-bg: {card_bg};
        --text-color: {text_color};
        --border-color: {border_color};
        --info-bg: {info_bg};
        --info-border: {info_border};
    }}

    /* Main container styling */
    .stApp {{
        background-color: var(--background-color);
        color: var(--text-color);
    }}
    
    /* Content container */
    .main .block-container {{
        padding: 2rem;
        max-width: 1200px;
    }}
    
    /* Header styling */
    .header-container {{
        background: var(--card-bg);
        padding: 2rem;
        border-radius: 15px;
        box-shadow: 0 10px 30px rgba(0,0,0,0.1);
        margin-bottom: 2rem;
        text-align: center;
    }}
    
    .header-title {{
        font-size: 2.5rem;
        font-weight: 700;
        background: linear-gradient(135deg, {header_gradient_start} 0%, {header_gradient_end} 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 0.5rem;
    }}
    
    .header-subtitle {{
        color: var(--text-color);
        opacity: 0.8;
        font-size: 1.1rem;
    }}
    
    /* Card styling - used for general containers like sidebar elements */
    .stSidebar .stMarkdown, .stSidebar .stSlider, .stSidebar .stMetric {{
        color: var(--text-color);
    }}
    
    /* Button styling */
    .stButton>button {{
        width: 100%;
        background: var(--primary-color);
        color: white;
        border: none;
        padding: 0.75rem 2rem;
        border-radius: 8px;
        font-weight: 600;
        font-size: 1rem;
        transition: transform 0.2s, box-shadow 0.2s;
    }}
    
    .stButton>button:hover {{
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(var(--primary-color), 0.4);
    }}
    
    /* Download button styling */
    .stDownloadButton>button {{
        width: 100%;
        background: var(--secondary-color);
        color: white;
        border: none;
        padding: 0.75rem 2rem;
        border-radius: 8px;
        font-weight: 600;
        font-size: 1rem;
        transition: transform 0.2s, box-shadow 0.2s;
    }}
    
    .stDownloadButton>button:hover {{
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(var(--secondary-color), 0.4);
    }}
    
    /* Text area styling */
    .stTextArea textarea {{
        border-radius: 8px;
        border: 2px solid var(--border-color);
        font-size: 1rem;
        background-color: var(--card-bg);
        color: var(--text-color);
    }}
    
    .stTextArea textarea:focus {{
        border-color: var(--primary-color);
        box-shadow: 0 0 0 2px rgba(var(--primary-color), 0.2);
    }}
    
    /* File uploader styling */
    .stFileUploader {{
        background: var(--card-bg);
        padding: 1rem;
        border-radius: 8px;
        border: 2px dashed var(--primary-color);
    }}
    
    /* Info box */
    .info-box {{
        background: var(--info-bg);
        color: var(--text-color);
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid var(--info-border);
        margin: 1rem 0;
    }}
    
    /* Streamlit components color adjustments for dark mode */
    .stTextInput>div>div>input, .stSelectbox>div>div>div>input, .stRadio>div>label, .stCheckbox>label, .stFileUploader>div>div>label {{
        color: var(--text-color);
    }}
    
    /* Adjust Streamlit's default elements that are hard to target */
    .st-emotion-cache-1cypcdb, .st-emotion-cache-10trblm, .st-emotion-cache-1v0mbdj {{
        background-color: var(--card-bg) !important;
        color: var(--text-color) !important;
    }}
    
    /* Sidebar background */
    .st-emotion-cache-vk3y59 {{ /* Target the sidebar container */
        background-color: var(--card-bg) !important;
    }}
    
    /* Sidebar header */
    .st-emotion-cache-16txte {{
        color: var(--text-color) !important;
    }}
    
    /* Adjust for Streamlit's internal theming on text/markdown */
    h1, h2, h3, h4, h5, h6, p, .stMarkdown {{
        color: var(--text-color);
    }}
    
    </style>
    """
    return css

st.markdown(get_custom_css(st.session_state.dark_mode), unsafe_allow_html=True)
# --- End Custom CSS ---


def create_docx_from_text(text):
    """
    Create a DOCX document from text, preserving paragraph structure (newlines).
    
    Args:
        text: str - The text to convert to DOCX
        
    Returns:
        BytesIO: Buffer containing the DOCX file
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Split text by paragraphs and add them to document
    paragraphs = text.split('\n')
    
    for para_text in paragraphs:
        # Add paragraph. The text area input preserves newlines, so we treat each line as a paragraph.
        # This is the best we can do for preserving structure from a plain text input.
        doc.add_paragraph(para_text)
    
    # Save to BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def save_docx_to_output(buffer, filename):
    """
    Save DOCX buffer to output folder.
    
    Args:
        buffer: BytesIO - The DOCX file buffer
        filename: str - Base filename (without extension)
    """
    try:
        # Create output directory if it doesn't exist
        output_dir = "output"
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate output filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = os.path.join(output_dir, f"{filename}_humanized_{timestamp}.docx")
        
        # Save document
        with open(output_path, "wb") as f:
            f.write(buffer.getbuffer())
        
        return output_path
    except Exception as e:
        st.error(f"❌ Error saving DOCX: {str(e)}")
        return None

def process_text_chunks(text, driver, chunk_size):
    """
    Splits text into chunks, humanizes each chunk, and returns the combined result.
    """
    final_humanized_text = ""
    chunks = split_text_preserve_paragraphs_and_newlines(text, chunk_size)
    
    st.info(f"Text split into {len(chunks)} chunks for processing.")
    
    for i, chunk in enumerate(chunks, 1):
        st.info(f"Processing Chunk {i}/{len(chunks)}...")
        
        try:
            result = get_texttohuman_humanizer_final(chunk, driver, save_debug=False)
            if result:
                # Use a single newline to join chunks, as the chunk content already contains internal newlines
                final_humanized_text += result + "\n"
            else:
                st.warning(f"Chunk {i} returned no result. Skipping.")
                
        except Exception as e:
            st.error(f"Error processing chunk {i}: {e}")
            # Continue with next chunk instead of failing completely
            continue
            
    return final_humanized_text.strip()

def handle_process_click():
    """
    Handles the main processing logic when the button is clicked.
    """
    st.session_state.processing = True
    st.session_state.humanized_text = ""
    st.session_state.docx_buffer = None
    
    if st.session_state.input_method == "Upload DOCX":
        # DOCX flow: The humanization is done inside read_docx_and_humanize
        if 'uploaded_file_path' not in st.session_state or not st.session_state.uploaded_file_path:
            st.error("Please upload a DOCX file first.")
            st.session_state.processing = False
            return
        
        with st.spinner("Initializing browser and humanizing DOCX... This may take a moment."):
            try:
                # Initialize driver (PlaywrightHumanizer context manager handles it)
                from texttohuman import PlaywrightHumanizer
                with PlaywrightHumanizer(headless=True, debug=False) as driver:
                    st.session_state.docx_buffer = read_docx_and_humanize(
                        st.session_state.uploaded_file_path, 
                        driver, 
                        chunk_size=st.session_state.chunk_size
                    )
                
                if st.session_state.docx_buffer:
                    st.success("✅ DOCX Humanization Complete!")
                    
                    # Read the text from the humanized DOCX for display in the output text area
                    # Note: This text extraction loses formatting, but is necessary for the Streamlit text area display.
                    humanized_doc = Document(st.session_state.docx_buffer)
                    humanized_text = "\n".join([p.text for p in humanized_doc.paragraphs])
                    st.session_state.humanized_text = humanized_text
                    
                    # Save the DOCX file to the output folder
                    st.session_state.output_filename = save_docx_to_output(
                        st.session_state.docx_buffer, 
                        st.session_state.input_filename
                    )
                else:
                    st.error("❌ DOCX Humanization Failed. Check logs for details.")
                    
            except Exception as e:
                st.error(f"An unexpected error occurred during DOCX processing: {e}")
            finally:
                # Clean up temp file
                if 'uploaded_file_path' in st.session_state and os.path.exists(st.session_state.uploaded_file_path):
                    os.remove(st.session_state.uploaded_file_path)
                    del st.session_state.uploaded_file_path
                st.session_state.processing = False
                st.rerun()
                
    else:
        # Text input flow
        input_text = st.session_state.text_input
        if not input_text.strip():
            st.error("Please enter some text to humanize.")
            st.session_state.processing = False
            return
            
        with st.spinner("Initializing browser and humanizing text... This may take a moment."):
            try:
                # Initialize driver (PlaywrightHumanizer context manager handles it)
                from texttohuman import PlaywrightHumanizer
                with PlaywrightHumanizer(headless=True, debug=False) as driver:
                    humanized_text = process_text_chunks(
                        input_text, 
                        driver, 
                        st.session_state.chunk_size
                    )
                
                if humanized_text:
                    st.session_state.humanized_text = humanized_text
                    st.success("✅ Text Humanization Complete!")
                    
                    # Create DOCX buffer for download and saving
                    st.session_state.docx_buffer = create_docx_from_text(humanized_text)
                    
                    # Save the DOCX file to the output folder
                    st.session_state.output_filename = save_docx_to_output(
                        st.session_state.docx_buffer, 
                        st.session_state.input_filename
                    )
                else:
                    st.error("❌ Text Humanization Failed. Check logs for details.")
                    
            except Exception as e:
                st.error(f"An unexpected error occurred during text processing: {e}")
            finally:
                st.session_state.processing = False
                st.rerun()

# Header
st.markdown("""
    <div class="header-container">
        <h1 class="header-title">✨ autohumanize-app : AI Text Humanizer</h1>
        <p class="header-subtitle">Transform AI-generated text into natural, human-like content</p>
    </div>
""", unsafe_allow_html=True)

# Sidebar
with st.sidebar:
    st.markdown("### ⚙️ Settings")
    
    # Dark Mode Toggle
    st.session_state.dark_mode = st.toggle(
        "🌙 Enable Dark Mode",
        value=st.session_state.dark_mode,
        key="dark_mode_toggle"
    )
    
    chunk_size = st.slider(
        "Chunk Size (words)",
        min_value=500,
        max_value=3000,
        value=2000,
        step=100,
        key="chunk_size",
        help="Split long texts into chunks of this size"
    )
    
    st.markdown("---")
    st.markdown("### 📊 Statistics")
    if st.session_state.humanized_text:
        word_count = len(st.session_state.humanized_text.split())
        char_count = len(st.session_state.humanized_text)
        st.metric("Words", word_count)
        st.metric("Characters", char_count)
        
        # Show saved file info
        if st.session_state.output_filename:
            st.markdown("---")
            st.markdown("### 💾 Saved File")
            st.success(f"📄 {os.path.basename(st.session_state.output_filename)}")
    else:
        st.info("Process text to see statistics")
    
    st.markdown("---")
    st.markdown("### ℹ️ About")
    st.markdown("""
    This tool uses advanced AI detection and humanization 
    to make your text sound more natural and human-written.
    
    **Features:**
    - Text & DOCX support
    - Auto-save to DOCX
    - Preserves formatting
    - Smart chunk processing
    - Copy to clipboard
    - Download as TXT/DOCX
    """)

# Main content area
col1, col2 = st.columns(2)

with col1:
    st.markdown("### 📝 Input")
    
    # Tab selection for input method
    input_method = st.radio(
        "Choose input method:",
        ["Type/Paste Text", "Upload DOCX"],
        horizontal=True,
        key="input_method"
    )
    
    input_text = ""
    
    if st.session_state.input_method == "Type/Paste Text":
        input_text = st.text_area(
            "Enter your text here:",
            height=400,
            placeholder="Paste or type your AI-generated text here...",
            key="text_input"
        )
        st.session_state.input_filename = "manual_input"
        # Clear DOCX related state
        if 'uploaded_file_path' in st.session_state:
            del st.session_state.uploaded_file_path
    else:
        uploaded_file = st.file_uploader(
            "Upload a DOCX file",
            type=['docx'],
            help="Upload a Word document to humanize"
        )
        
        if uploaded_file is not None:
            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                st.session_state.uploaded_file_path = tmp_file.name
            
            # Store original filename (without extension)
            st.session_state.input_filename = os.path.splitext(uploaded_file.name)[0]
            
            st.success(f"File uploaded: {uploaded_file.name}")
        else:
            # Clear file path if file is removed
            if 'uploaded_file_path' in st.session_state and os.path.exists(st.session_state.uploaded_file_path):
                os.remove(st.session_state.uploaded_file_path)
                del st.session_state.uploaded_file_path
            st.session_state.input_filename = ""
            
    st.button(
        "✨ Humanize Text",
        on_click=handle_process_click,
        disabled=st.session_state.processing,
        use_container_width=True
    )

with col2:
    st.markdown("### 📄 Output")
    
    st.text_area(
        "Humanized Text",
        st.session_state.humanized_text,
        height=400,
        key="output_text",
        disabled=True
    )
    
    col_dl1, col_dl2 = st.columns(2)
    
    if st.session_state.humanized_text:
        # Download as DOCX
        if st.session_state.docx_buffer:
            docx_filename = f"{st.session_state.input_filename}_humanized.docx"
            col_dl1.download_button(
                label="⬇️ Download DOCX",
                data=st.session_state.docx_buffer,
                file_name=docx_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        else:
            col_dl1.info("DOCX not available for download.")
            
        # Download as TXT
        txt_filename = f"{st.session_state.input_filename}_humanized.txt"
        col_dl2.download_button(
            label="⬇️ Download TXT",
            data=st.session_state.humanized_text.encode('utf-8'),
            file_name=txt_filename,
            mime="text/plain",
            use_container_width=True
        )
        
        # Copy to clipboard (requires custom JS/HTML, which is complex in Streamlit, so we'll skip for now or rely on the user to copy from the text area)
        # st.button("📋 Copy to Clipboard", use_container_width=True)
    else:
        col_dl1.info("Output will appear here.")
        col_dl2.info("Output will appear here.")
        
st.markdown("---")
st.markdown("""
    <div class="info-box">
        **Note on Formatting:** For DOCX uploads, the original document structure (headings, tables, bold text) is preserved by editing the document in place. For text input, newlines are treated as paragraph breaks to maintain structure.
    </div>
""", unsafe_allow_html=True)
